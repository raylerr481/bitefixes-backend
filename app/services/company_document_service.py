"""Business document ingestion for Company AI Profiles."""
from __future__ import annotations
import hashlib
import io
import re
from datetime import datetime, timezone
from typing import Any, Dict
from app.database.supabase import database
from app.services.company_service import get_company, get_company_ai_profile

def _clean_text(text: str) -> str:
    text = re.sub(r"\x00", " ", text or "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def _extract_text(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in {"txt", "csv", "json", "md"}:
        return _clean_text(content.decode("utf-8-sig", errors="replace"))
    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return _clean_text("\n\n".join(page.extract_text() or "" for page in reader.pages))
    if ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return _clean_text("\n".join(parts))
    if ext == "doc":
        raise ValueError("DOC legacy files are not supported; convert to DOCX or PDF.")
    raise ValueError("Unsupported document type")

def _upsert_profile(company_id: int, company_name: str, filename: str, text: str, sha256: str) -> Dict[str, Any]:
    now = datetime.now(timezone.utc).isoformat()
    existing = get_company_ai_profile(company_id) or {}
    profile = existing.get("profile") if isinstance(existing.get("profile"), dict) else {}
    documents = profile.get("documents") if isinstance(profile.get("documents"), list) else []
    documents = [d for d in documents if not (isinstance(d, dict) and d.get("sha256") == sha256)]
    documents.append({"name": filename, "sha256": sha256, "ingested_at": now, "characters": len(text)})
    profile["documents"] = documents[-50:]
    profile["document_context"] = text[:120000]
    profile["context_source"] = "wordpress_company_document"
    payload = {"company_id": company_id, "company_name": company_name or existing.get("company_name") or (get_company(company_id) or {}).get("name"), "description": existing.get("description") or "", "industry": existing.get("industry") or "", "profile": profile, "updated_at": now}
    if existing.get("id"):
        result = database.table("company_ai_profiles").update(payload).eq("id", existing["id"]).execute()
    else:
        result = database.table("company_ai_profiles").insert(payload).execute()
    if not result.data:
        raise RuntimeError("Company AI Profile could not be persisted")
    return result.data[0]

def ingest_company_document(*, company_id: int, company_name: str, filename: str, content: bytes, content_type: str) -> Dict[str, Any]:
    if not get_company(company_id):
        raise LookupError("Company not found")
    if not content:
        raise ValueError("Empty document")
    sha256 = hashlib.sha256(content).hexdigest()
    text = _extract_text(filename, content)
    if not text:
        raise ValueError("No readable text was extracted from the document")
    profile = _upsert_profile(company_id, company_name, filename, text, sha256)
    return {"status": "ingested", "company_id": company_id, "profile_id": profile.get("id"), "filename": filename, "content_type": content_type, "sha256": sha256, "characters_extracted": len(text), "context_updated": True, "external_ai_required": False, "external_ai_next_step": "available_for_contextual_reasoning"}
